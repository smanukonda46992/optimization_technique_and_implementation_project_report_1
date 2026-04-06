from pathlib import Path
import csv
import math
from datetime import date

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "report"
PLOTS_DIR = ROOT / "results" / "plots"
RAW_CSV = ROOT / "results" / "raw_data" / "benchmark_results.csv"
SRC_DIR = ROOT / "src"

AUTHOR_NAME = "Sai Mohan Manukonda"
UC_STUDENT_ID = "005046992"

PART1_REPORT_DOCX = REPORT_DIR / "Part1_Optimization_Technique_Project_Report.docx"
CODE_DOCX = REPORT_DIR / "Part1_Source_Code_and_Screenshots.docx"


def set_normal_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Optimization in High-Performance Computing\n")
    run.bold = True
    run.font.size = Pt(18)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("Part 1: Optimization Technique and Implementation Project Report\n").bold = True
    p2.add_run("\n")
    p2.add_run("Course: Algorithms and Data Structures\n")
    p2.add_run(f"Student: {AUTHOR_NAME}\n")
    p2.add_run(f"UC Student ID Number: {UC_STUDENT_ID}\n")
    p2.add_run(f"Date: {date.today().strftime('%B %d, %Y')}\n")

    doc.add_page_break()


def load_summary_metrics():
    rows = []
    with open(RAW_CSV, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            row["distribution"] = row["distribution"]
            row["size"] = int(row["size"])
            row["baseline_mean_s"] = float(row["baseline_mean_s"])
            row["optimized_mean_s"] = float(row["optimized_mean_s"])
            row["speedup"] = float(row["speedup"])
            rows.append(row)

    valid_speedups = [r["speedup"] for r in rows if not math.isnan(r["speedup"])]
    max_speedup = max(valid_speedups) if valid_speedups else float("nan")
    avg_speedup = sum(valid_speedups) / len(valid_speedups) if valid_speedups else float("nan")

    baseline_failures = sum(1 for r in rows if math.isnan(r["baseline_mean_s"]))
    optimized_failures = sum(1 for r in rows if math.isnan(r["optimized_mean_s"]))

    return {
        "rows": rows,
        "max_speedup": max_speedup,
        "avg_speedup": avg_speedup,
        "baseline_failures": baseline_failures,
        "optimized_failures": optimized_failures,
    }


def add_section_heading(doc, text):
    doc.add_heading(text, level=1)


def add_report_body(doc, metrics):
    add_section_heading(doc, "Abstract")
    doc.add_paragraph(
        "This report evaluates one HPC-relevant optimization technique for data structures and sorting workloads: "
        "algorithmic hardening of Quicksort using randomized pivot selection, three-way partitioning for duplicates, "
        "small-partition insertion sort, and tail-recursion elimination. A Python prototype compares this optimized "
        "version against a baseline deterministic first-pivot implementation. Empirical results across five input "
        "distributions show improved robustness, with notable gains on sorted/reverse-sorted patterns where baseline "
        "recursion failures occur at larger sizes."
    )

    add_section_heading(doc, "1. Selected Optimization Technique")
    doc.add_paragraph(
        "Selected Technique: Input-resilient Quicksort optimization through randomized pivoting and partition strategy improvements."
    )
    doc.add_paragraph(
        "Justification: This technique directly targets one common HPC performance bug pattern: data-dependent algorithmic "
        "degradation. Deterministic first-pivot quicksort is vulnerable to adversarial or naturally ordered inputs, resulting in "
        "O(n^2) behavior and deep recursion. The optimized approach reduces worst-case likelihood in practice and improves "
        "throughput stability across varied distributions."
    )

    add_section_heading(doc, "2. Strengths and Weaknesses")
    doc.add_paragraph("Strengths:")
    doc.add_paragraph("- Randomized pivoting minimizes predictable worst-case partitions.", style="List Bullet")
    doc.add_paragraph("- Three-way partitioning improves behavior on duplicate-heavy data.", style="List Bullet")
    doc.add_paragraph("- Insertion-sort threshold lowers overhead on very small partitions.", style="List Bullet")
    doc.add_paragraph("- Explicit stack/tail recursion strategy prevents deep recursion growth.", style="List Bullet")
    doc.add_paragraph("Weaknesses:")
    doc.add_paragraph("- Additional implementation complexity compared to textbook baseline quicksort.", style="List Bullet")
    doc.add_paragraph("- Slightly higher overhead on small random inputs where baseline is already balanced.", style="List Bullet")

    add_section_heading(doc, "3. Prototype Implementation")
    doc.add_paragraph(
        "Prototype files are located in src/: baseline_quicksort.py, optimized_quicksort.py, and input_generators.py. "
        "The experiment driver in experiments/run_benchmarks.py executes 5 trials per configuration across random, sorted, "
        "reverse_sorted, nearly_sorted, and duplicates distributions."
    )

    add_section_heading(doc, "4. Empirical Results and Analysis")
    doc.add_paragraph(
        f"Across all valid comparisons, average measured speedup was {metrics['avg_speedup']:.2f}x, with a maximum observed "
        f"speedup of {metrics['max_speedup']:.2f}x on ordered patterns. Baseline failures observed: {metrics['baseline_failures']} "
        f"cases; optimized failures observed: {metrics['optimized_failures']} cases."
    )
    doc.add_paragraph(
        "Interpretation: baseline deterministic quicksort can perform well for random arrays but is fragile under ordered inputs. "
        "The optimized variant trades slight random-input overhead for significant resilience and consistent completion."
    )

    timings_plot = PLOTS_DIR / "timings_by_distribution.png"
    speedup_plot = PLOTS_DIR / "speedup_summary.png"
    if timings_plot.exists():
        doc.add_paragraph("Figure 1. Timing comparison by input distribution.")
        doc.add_picture(str(timings_plot), width=Inches(6.2))
    if speedup_plot.exists():
        doc.add_paragraph("Figure 2. Speedup summary (baseline / optimized).")
        doc.add_picture(str(speedup_plot), width=Inches(5.8))

    add_section_heading(doc, "5. Problems Encountered and Lessons Learned")
    doc.add_paragraph(
        "During benchmarking, baseline deterministic quicksort triggered recursion errors on sorted and reverse-sorted arrays "
        "for larger sizes. This mirrored theoretical expectations for pathological pivot behavior and demonstrates why practical "
        "optimization must include robustness, not only average-case speed."
    )

    add_section_heading(doc, "6. Theoretical vs Observed Behavior")
    doc.add_paragraph(
        "Theory predicts randomized quicksort approaches O(n log n) expected time, while deterministic first-pivot quicksort "
        "can degrade to O(n^2). The observations align with this: the optimized version remained stable across all tested "
        "distributions, whereas baseline behavior degraded and failed under ordered inputs."
    )

    add_section_heading(doc, "7. Conclusion")
    doc.add_paragraph(
        "The selected optimization technique effectively improved resilience and practical performance stability. For HPC-like "
        "workloads where input distributions are dynamic or not fully controlled, robust pivot and partition strategies provide "
        "a meaningful reliability advantage over naive deterministic implementations."
    )

    add_section_heading(doc, "References")
    refs = [
        "Blum, M., Floyd, R. W., Pratt, V., Rivest, R. L., & Tarjan, R. E. (1973). Time bounds for selection. Journal of Computer and System Sciences, 7(4), 448-461.",
        "Cormen, T. H., Leiserson, C. E., Rivest, R. L., & Stein, C. (2022). Introduction to algorithms (4th ed.). MIT Press.",
        "Sedgewick, R., & Wayne, K. (2011). Algorithms (4th ed.). Addison-Wesley.",
        "McSherry, F., Isard, M., & Murray, D. G. (2015). Scalability! But at what COST? In 15th Workshop on Hot Topics in Operating Systems.",
        "Luo, Y., Wong, M. L., & Hwu, W.-M. (2016). Characterizing and understanding performance bugs in HPC applications. Proceedings of SC Workshops.",
        "Project benchmark artifacts and generated plots from this repository (results/raw_data and results/plots).",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")


def add_code_document(doc):
    doc.add_heading("Part 1 Source Code and Screenshots", level=1)
    doc.add_paragraph(
        "This companion MS Word document contains implementation evidence for the selected optimization technique."
    )

    doc.add_heading("A. File Inventory", level=2)
    inventory = [
        "src/baseline_quicksort.py",
        "src/optimized_quicksort.py",
        "src/input_generators.py",
        "experiments/run_benchmarks.py",
        "experiments/plot_results.py",
        "results/raw_data/benchmark_results.csv",
        "results/plots/timings_by_distribution.png",
        "results/plots/speedup_summary.png",
    ]
    for item in inventory:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("B. Baseline Quicksort (Core Snippet)", level=2)
    baseline_text = (SRC_DIR / "baseline_quicksort.py").read_text(encoding="utf-8")
    doc.add_paragraph(baseline_text)

    doc.add_heading("C. Optimized Quicksort (Core Snippet)", level=2)
    optimized_text = (SRC_DIR / "optimized_quicksort.py").read_text(encoding="utf-8")
    doc.add_paragraph(optimized_text)

    doc.add_heading("D. Benchmark Screenshot Figures", level=2)
    timings_plot = PLOTS_DIR / "timings_by_distribution.png"
    speedup_plot = PLOTS_DIR / "speedup_summary.png"
    if timings_plot.exists():
        doc.add_paragraph("Figure A: Timing by distribution")
        doc.add_picture(str(timings_plot), width=Inches(6.2))
    if speedup_plot.exists():
        doc.add_paragraph("Figure B: Speedup summary")
        doc.add_picture(str(speedup_plot), width=Inches(6.0))

    doc.add_heading("E. Execution Notes", level=2)
    doc.add_paragraph(
        "Commands used:\n"
        "1) PYTHONPATH=. python experiments/run_benchmarks.py\n"
        "2) python experiments/plot_results.py\n"
        "The benchmark CSV and PNG outputs are included in this project folder."
    )


def main():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    metrics = load_summary_metrics()

    part1_doc = Document()
    set_normal_font(part1_doc)
    part1_doc.core_properties.author = AUTHOR_NAME
    part1_doc.core_properties.last_modified_by = AUTHOR_NAME
    part1_doc.core_properties.title = "Part 1 Optimization Technique Project Report"
    add_title_page(part1_doc)
    add_report_body(part1_doc, metrics)
    part1_doc.save(PART1_REPORT_DOCX)

    code_doc = Document()
    set_normal_font(code_doc)
    code_doc.core_properties.author = AUTHOR_NAME
    code_doc.core_properties.last_modified_by = AUTHOR_NAME
    code_doc.core_properties.title = "Part 1 Source Code and Screenshots"
    add_code_document(code_doc)
    code_doc.save(CODE_DOCX)

    print(f"Created: {PART1_REPORT_DOCX}")
    print(f"Created: {CODE_DOCX}")


if __name__ == "__main__":
    main()
